import streamlit as st
import joblib
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from sklearn.cluster import DBSCAN
from appsynthetic_signals import generate_synthetic_ecg
from appsignal_preprocessing import bandpass_filter_custom
from sklearn.metrics import silhouette_score
from tslearn.preprocessing import TimeSeriesScalerMeanVariance
from scipy.signal import resample
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from ecg_db import (
    register_user,
    authenticate_user,
    save_result,
    get_results,
    update_result,
    delete_result,
    init_db,
)


# Custom CSS to style the app
st.markdown("""
    <style>
    /* Change sidebar background to black */
    .css-1d391kg {
        background-color: black;
        color: white;
    }

    /* Change the text color inside the sidebar */
    .css-1d391kg a, .css-1d391kg .st-ae, .css-1d391kg .st-cw {
        color: white;
    }

    /* Style for the sidebar navigation items */
    .css-1d391kg .st-ae:hover {
        color: #f0f0f0; /* Lighter color on hover */
    }

    /* Make sure all text in sidebar remains white */
    .st-sidebar .sidebar-content {
        color: white;
    }

    /* Style for the main content */
    .main {
        background-color: white;
    }

    /* Custom style for the header */
    .css-1k4rji1 h1 {
        font-size: 2em;
    }

    /* Custom button styling */
    .stButton > button {
        background-color: #007bff;
        color: white;
        border: none;
        padding: 10px 20px;
        font-size: 16px;
        border-radius: 5px;
    }
    .stButton > button:hover {
        background-color: #0056b3;
    }

    /* Override other page elements if needed */
    .css-ffhzg2 {
        background-color: white;
    }
    </style>
    """, unsafe_allow_html=True)

init_db()

# Load pre-trained k-Shape and DBSCAN models
@st.cache_resource
def load_kshape_model():
    return joblib.load("C:/Users/user/ai_env/ICSproject2/outputs/clusters and metrics/kshape_model.pkl")  # Adjust path if needed

@st.cache_resource
def load_dbscan_model():
    return joblib.load("C:/Users/user/ai_env/ICSproject2/outputs/clusters and metrics/dbscan_model.pkl")  # Adjust path if needed

# Helper functions
def preprocess_signal(signal, target_length=250):
    """Resample and reshape signal to match k-Shape input dimensions"""
    resampled_signal = resample(signal, target_length)
    reshaped_signal = resampled_signal.reshape(1, target_length, 1)
    reshaped_signal = np.repeat(reshaped_signal, 3, axis=0)  # Repeat 3 times to match (3, target_length, 1)
    return reshaped_signal

# Function to generate a Word report
def generate_word_report(initial_signal, processed_signal, cluster_id, dbscan_labels, dbscan_plot):
    document = Document()
    document.add_heading("ECG Signal Analysis Report", level=1)
    
    # Initial Signal Details
    document.add_heading("Initial Signal", level=2)
    document.add_paragraph(f"Total Samples: {len(initial_signal)}")
    document.add_paragraph(f"First 10 values: {initial_signal[:10]}")
    
    # Processed Signal Details
    document.add_heading("Processed Signal", level=2)
    document.add_paragraph(f"First 10 values: {processed_signal[:10]}")
    
    # Clustering Results
    document.add_heading("Clustering Results", level=2)
    document.add_paragraph(f"k-Shape Cluster ID: {cluster_id}")
    document.add_paragraph(f"DBSCAN Labels (First 10): {dbscan_labels[:10]}")
    
    # DBSCAN Plot
    document.add_heading("DBSCAN Clustering Plot", level=2)
    img_stream = BytesIO()
    dbscan_plot.savefig(img_stream, format="png")
    img_stream.seek(0)
    document.add_picture(img_stream, width=5)  # Adjust width to fit page
    
    return document

# Function to generate a PDF report
def generate_pdf_report(initial_signal, processed_signal, cluster_id, dbscan_labels, dbscan_plot):
    pdf_stream = BytesIO()
    c = canvas.Canvas(pdf_stream, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, 750, "ECG Signal Analysis Report")
    
    # Initial Signal Details
    c.setFont("Helvetica", 12)
    c.drawString(50, 720, f"Initial Signal: First 10 values: {initial_signal[:10]}")
    
    # Processed Signal Details
    c.drawString(50, 700, f"Processed Signal: First 10 values: {processed_signal[:10]}")
    
    # Clustering Results
    c.drawString(50, 680, f"k-Shape Cluster ID: {cluster_id}")
    c.drawString(50, 660, f"DBSCAN Labels (First 10): {dbscan_labels[:10]}")
    
    # DBSCAN Plot
    img_stream = BytesIO()
    dbscan_plot.savefig(img_stream, format="png")
    img_stream.seek(0)
    c.drawInlineImage(img_stream, 50, 400, width=400, height=200)
    
    c.save()
    pdf_stream.seek(0)
    
    return pdf_stream

# Function to show the registration form
def show_registration_form():
    st.title("ECG Signal Pattern Recognition")
    st.write("Register")

    # Username and Password fields
    username = st.text_input("Enter a Username")
    password = st.text_input("Enter a Password", type="password")
    
    # Registration button
    if st.button("Register"):
        if username and password:
            # Register user and handle potential errors
            if register_user(username, password):
                st.success("Registration successful! You can now log in.")
            else:
                st.error("Username already exists. Please try a different one.")
        else:
            st.warning("Please fill in both username and password.")
    
    # Button to switch to login form
    if st.button("Already have an account? Log in"):
        st.session_state.page = "login"  # Change the page to login

# Function to show the login form
def show_login_form():
    st.title("ECG Signal Pattern Recognition")
    st.write("Login")

    # Username and Password fields
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    # Login button
    if st.button("Login"):
        if username and password:
            user = authenticate_user(username, password)
            if user:
                st.session_state.logged_in = True
                st.session_state.user_id = user[0]  # Store user ID
                st.success("Login Successful!")
                st.rerun()  # Reload the page to reflect logged-in state
            else:
                st.error("Invalid username or password.")
        else:
            st.warning("Please fill in both username and password.")

# Check which page to show based on the session state
if "logged_in" not in st.session_state or not st.session_state.logged_in:
    if "page" not in st.session_state or st.session_state.page == "register":
        # Show the registration form
        show_registration_form()
    elif st.session_state.page == "login":
        # Show the login form
        show_login_form()
else:
    # Once logged in, display the home page


    # Streamlit UI
    st.title("ECG Signal Analysis")
    st.write("Navigate the steps using the options on the sidebars")

    # Navigation Sidebar
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Select a page", ["Signal Input", "Preprocessing", "Pattern Recognition", "Reports"], key="page_selector")

    # Signal Input
    if page == "Signal Input":
        st.write("Input signal here")
        signal_input_option = st.selectbox("Choose your signal input:", ("Upload ECG Signal", "Generate Synthetic ECG"))

        if signal_input_option == "Upload ECG Signal":
            uploaded_file = st.file_uploader("Upload ECG Signal File", type=["csv", "txt", "json"])
            
            if uploaded_file is not None:
                try:
                    uploaded_data = np.loadtxt(uploaded_file)
                    st.line_chart(uploaded_data)
                    st.session_state.signal = uploaded_data  # Store signal in session state
                    st.session_state.sampling_rate = st.sidebar.number_input(
                        "Enter Sampling Rate (Hz) for Uploaded Signal", 
                        min_value=100, 
                        max_value=5000, 
                        value=1000, 
                        step=100,
                        key="uploaded_sampling_rate"
                    )
                except Exception as e:
                    st.error(f"Error reading the uploaded file: {e}")
        elif signal_input_option == "Generate Synthetic ECG":
            hr = st.slider("Select Heart Rate (BPM)", min_value=40, max_value=120, value=75, key="hr_slider")
            noise_level = st.slider("Select Noise Level", min_value=0.0, max_value=0.5, value=0.05, key="noise_slider")
            length = st.slider("Select Signal Length (samples)", min_value=1000, max_value=10000, value=5000, key="length_slider")
            sampling_rate = st.sidebar.slider("Sampling Rate (Hz)", min_value=500, max_value=5000, value=1000, step=100, key="sampling_rate_slider")
            
            signal = generate_synthetic_ecg(hr=hr, noise_level=noise_level, length=length, sampling_rate=sampling_rate)
            st.line_chart(signal)
            st.session_state.signal = signal  # Store signal in session state
            st.session_state.sampling_rate = sampling_rate

    # Convert numpy array to CSV string
    def convert_to_csv(arr):
        df = pd.DataFrame(arr)
        csv = df.to_csv(index=False)
        return csv

    # Preprocessing section
    if page == "Preprocessing":
        if "signal" not in st.session_state:
            st.warning("Please upload or generate a signal first.")
        else:
            st.sidebar.header("Signal Preprocessing")
            st.write("Preprocess signal by adjusting lowcut and highcut values using sliders on the sidebar.")
        
            lowcut = st.sidebar.slider("Low Cutoff Frequency (Hz)", min_value=0.1, max_value=5.0, value=0.5, step=0.1, key="lowcut_slider")
            highcut = st.sidebar.slider("High Cutoff Frequency (Hz)", min_value=5.0, max_value=50.0, value=40.0, step=1.0, key="highcut_slider")

            signal = st.session_state.signal
            sampling_rate = st.session_state.sampling_rate

            # Apply bandpass filter
            processed_signal = bandpass_filter_custom(signal, lowcut=lowcut, highcut=highcut, fs=sampling_rate)
            st.subheader("Processed Signal")
            st.line_chart(processed_signal)

            # Convert to CSV for download
            csv_signal = convert_to_csv(processed_signal)

            # Provide download button
            st.download_button("Download Processed Signal", csv_signal, file_name="processed_signal.csv", mime="text/csv")
    if page == "Pattern Recognition":
    # Initialize session states
        if "dbscan_labels" not in st.session_state:
          st.session_state.dbscan_labels = None
        if "kshape_clusters" not in st.session_state:
          st.session_state.kshape_clusters = None
        if "processed_signal" not in st.session_state:
          st.session_state.processed_signal = None

        if "signal" not in st.session_state:
           st.warning("Please upload or generate a signal first.")
        else:
           if st.button("Run Pattern Recognition"):
             signal = st.session_state.signal
             sampling_rate = st.session_state.sampling_rate

            # Load the pre-trained models for k-Shape and DBSCAN
             kshape_model = load_kshape_model()
             dbscan_model = load_dbscan_model()

            # Preprocess the signal before applying clustering (bandpass filter)
             processed_signal = bandpass_filter_custom(signal, lowcut=0.5, highcut=40.0, fs=sampling_rate)

            # Reshape signal for DBSCAN (flatten the time series for clustering)
             processed_signal_flat = processed_signal.reshape(-1, 1)

            # Apply DBSCAN
             eps = 0.5  # Example value for DBSCAN
             min_samples = 5  # Example value for DBSCAN
             dbscan = DBSCAN(eps=eps, min_samples=min_samples)
             dbscan_labels = dbscan.fit_predict(processed_signal_flat)

            # Apply k-Shape clustering
             processed_signal_reshaped = preprocess_signal(processed_signal)
             kshape_clusters = kshape_model.predict(processed_signal_reshaped)

            # Save results to session state
             st.session_state.processed_signal = processed_signal
             st.session_state.dbscan_labels = dbscan_labels
             st.session_state.kshape_clusters = kshape_clusters

             st.success("Clustering completed!.")

            # Display results
             col1, col2 = st.columns(2)

             with col1:
                st.subheader("DBSCAN Clustering")
                plt.figure(figsize=(8, 6))
                plt.plot(processed_signal_flat, label="Signal", color='grey', alpha=0.5)
                plt.scatter(np.arange(len(processed_signal_flat)), processed_signal_flat, c=dbscan_labels, cmap="viridis", s=5)
                plt.title("DBSCAN Clustering")
                plt.xlabel("Sample Index")
                plt.ylabel("Amplitude")
                plt.legend()
                st.pyplot(plt)
                noise_points = np.sum(dbscan_labels == -1)
                st.write(f"Number of noise points detected by DBSCAN: {noise_points}")

             with col2:
                st.subheader("k-Shape Clustering")
                custom_cluster_names = {0: "Normal", 1: "Arrhythmia", 2: "Abnormal"}
                plt.figure(figsize=(8, 6))
                for i in range(np.max(kshape_clusters) + 1):
                    cluster_indices = np.where(kshape_clusters == i)[0]
                    cluster_name = custom_cluster_names.get(i, f"Cluster {i + 1}")
                    plt.plot(processed_signal, color='lightgray', alpha=0.5)
                    plt.scatter(cluster_indices, processed_signal[cluster_indices], label=cluster_name, s=10)
                plt.title("k-Shape Clustering")
                plt.xlabel("Sample Index")
                plt.ylabel("Amplitude")
                plt.legend()
                st.pyplot(plt)
                st.write(f"Cluster Assigned: {cluster_name}")

        st.subheader("Save Your Clustering Results")
        result_name = st.text_input("Enter a name for the result")

        if st.button("Save Clustering Results"):
            if result_name and st.session_state.dbscan_labels is not None and st.session_state.kshape_clusters is not None:
                # Calculate number of noise points
                noise_points = np.sum(st.session_state.dbscan_labels == -1)
                custom_cluster_names = {0: "Normal", 1: "Arrhythmia", 2: "Abnormal"}
                # Determine the cluster name based on the maximum cluster ID in k-Shape
                if len(st.session_state.kshape_clusters) > 0:  # Ensure kshape_clusters has data
                    max_cluster_id = int(np.max(st.session_state.kshape_clusters))  # Get max cluster ID
                    cluster_name = custom_cluster_names.get(max_cluster_id, f"Cluster {max_cluster_id}")
                else:
                    cluster_name = "Unknown Cluster"  # Fallback if kshape_clusters is empty
                # Prepare the result data to save
                result_data =  {
                     "noise_points":int(noise_points) ,
                      
                     "cluster_name": cluster_name
                    }
                save_result(st.session_state.user_id, result_name, str(result_data))
                st.success("Result saved successfully!")
            else:
                st.warning("No clustering results found or name not provided.")

   # View Saved Results Page
    if page == "Reports":
      st.header("View Saved Results")

# Button to fetch saved results
      if st.button("View Saved Results"):
         if "user_id" in st.session_state:  # Ensure user_id exists in session state
    # Fetch results for the logged-in user
             all_results = get_results(st.session_state.user_id)

    # Filter results for the logged-in user
             user_id = st.session_state.user_id
             user_results = [result for result in all_results if result[1] == user_id]  # Assuming user_id is the second field

             if user_results:  # Check if user-specific results are not empty
        # Create a DataFrame for user-specific results
                 
                 df = pd.DataFrame(user_results, columns=["Result ID", "User ID", "Result Name", "Result Data"])

        # Display results in a custom table format
                 st.write("**Manage Results**")
                 for index, row in df.iterrows():
                     result_id = row["Result ID"]
                     user_id = row["User ID"]
                     result_name = row["Result Name"]
                     result_data = row["Result Data"]

                     cols = st.columns([4,4, 1, 1])  # Adjust column widths as needed
                     #with cols[0]:
                        # st.write("Results ID")
                        # st.write(result_id)
                    # with cols[1]:
                         #st.write("Results ID")
                         #st.write(user_id)
                     with cols[0]:
                         st.write("Result Name:")
                         st.write(result_name)
                     with cols[1]:
                         st.write("Data:")
                         st.write(result_data)

            # Group actions inside a form to prevent page refreshing
                     with st.form(key=f"edit_delete_form_{result_id}", clear_on_submit=True):
                      for result_id, result_name, result_data in user_results:     
                         edit_button = st.form_submit_button("✏️ Edit")
                         delete_button = st.form_submit_button("🗑️ Delete")
        
             if edit_button:
                new_name = st.text_input(f"Edit Result Name for '{result_name}'", value=result_name, key=f"new_name_{result_id}")
                save_button = st.form_submit_button(f"Save Changes for {result_name}", key=f"save_{result_id}")
                if save_button:
                     if update_result(result_id, new_name, result_data):  # Keep result_data unchanged
                        st.session_state.edit_success = True
                        st.success(f"Result name updated to '{new_name}' successfully!")
                        st.experimental_rerun()
                     else:
                        st.error(f"Failed to update '{result_name}'.")

             if delete_button:
                     if delete_result(result_id):
                        st.session_state.delete_success = True
                        st.success(f"Result '{result_name}' deleted successfully!")
                        st.experimental_rerun()
                     else:
                        st.error(f"Failed to delete '{result_name}'.")

             else:
                 st.info(f"No saved results found for User {user_id}.")
         else:
             st.warning("User not logged in or session expired.")

